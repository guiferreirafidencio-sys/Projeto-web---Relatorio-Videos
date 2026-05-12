import whisper
import yt_dlp
from docx import Document
import google.generativeai as genai
from docx2pdf import convert
import os
import time
from dotenv import load_dotenv
import pythoncom
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import threading

# 🔐 carregar .env
load_dotenv()

status = ""
resultado_final = None

# 🔑 Lista de APIs
api_keys = [
    os.getenv("api_1"),
    os.getenv("api_2"),
    os.getenv("api_3"),
    os.getenv("api_4")
]

# 🔹 Prompt principal
prompt = """
Pegue o texto abaixo e faça as seguintes melhorias:

1. Corrija erros de português e frases mal estruturadas.
2. Mantenha o contexto original do texto.
3. Organize as informações de forma clara e lógica.
4. Estruture o conteúdo como um pequeno relatório.
5. Destaque os pontos mais importantes.
6. Use tópicos e subtítulos.

Texto:
"""

# 📥 Baixar áudio
def baixar_audio(url):
    global status

    status = "🔄 Baixando áudio..."

    # remove arquivos antigos
    arquivos = [
        "audio.mp3",
        "audio.webm",
        "audio.webm.part"
    ]

    for arquivo in arquivos:
        if os.path.exists(arquivo):
            try:
                os.remove(arquivo)
            except:
                pass

    opcoes = {
        'format': 'bestaudio/best',
        'outtmpl': 'audio.%(ext)s',
        'noplaylist': True,
        'quiet': True,
        'no_warnings': True,
        'nopart': True,  # evita erro .part
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'mp3',
        }],
    }

    with yt_dlp.YoutubeDL(opcoes) as ydl:
        ydl.download([url])

# 🎧 Transcrever
def transcrever():
    global status

    status = "🎧 Transcrevendo..."

    modelo = whisper.load_model("tiny")

    resultado = modelo.transcribe("audio.mp3")

    return resultado["text"]

# 🧠 Gerar resumo com troca automática de API
def gerar_resumo(texto):
    global status

    status = "🧠 Gerando resumo..."

    ultimo_erro = None

    for i, key in enumerate(api_keys):

        try:
            if not key:
                continue

            print(f"🔄 Tentando API {i+1}")

            genai.configure(api_key=key)

            modelo = genai.GenerativeModel("gemini-2.5-flash")

            resposta = modelo.generate_content(prompt + texto)

            print(f"✅ API {i+1} funcionando")

            return resposta.text

        except Exception as e:

            erro = str(e)
            ultimo_erro = erro

            print(f"❌ API {i+1} falhou")

            # se atingir limite/quota
            if (
                "429" in erro or
                "quota" in erro.lower() or
                "rate limit" in erro.lower()
            ):
                print("⚠️ Limite atingido. Tentando próxima API...")
                time.sleep(1)
                continue

            # outros erros
            print(f"⚠️ Outro erro: {erro}")
            time.sleep(1)
            continue

    raise Exception("Todas as APIs falharam.")

# 📄 Gerar DOCX + PDF
def gerar_doc(url, transcricao, resumo):
    global status

    status = "📄 Gerando documento..."

    nome_docx = "relatorio_video.docx"
    nome_pdf = "relatorio_video.pdf"

    doc = Document()

    titulo = doc.add_heading("📊 Relatório Analítico do Vídeo", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_paragraph("Gerado automaticamente por IA")
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    doc.add_heading("🔗 Fonte do Conteúdo", 1)
    doc.add_paragraph(url)

    doc.add_paragraph("—" * 40)

    doc.add_heading("🧠 Análise e Resumo Estruturado", 1)

    for linha in resumo.split("\n"):

        if linha.strip() == "":
            continue

        p = doc.add_paragraph(linha)

        if "##" in linha or "###" in linha:
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(14)

    doc.add_paragraph("")

    rodape = doc.add_paragraph(
        "⚡ Relatório gerado automaticamente • Guilherme Dev"
    )

    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(nome_docx)

    status = "📄 Convertendo para PDF..."

    convert(nome_docx, nome_pdf)

    return nome_pdf

# 🔥 PROCESSO PRINCIPAL
def processar(url):
    global status, resultado_final

    try:
        pythoncom.CoInitialize()

        status = "🚀 Iniciando processo..."

        baixar_audio(url)

        texto = transcrever()

        resumo = gerar_resumo(texto)

        pdf_final = gerar_doc(url, texto, resumo)

        resultado_final = resumo

        status = "✅ Finalizado"

        return pdf_final

    except Exception as e:

        print(f"ERRO INTERNO: {e}")

        # erro amigável frontend
        status = "❌ Ocorreu um erro no processamento"

        return None

    finally:
        pythoncom.CoUninitialize()

# 🚀 THREAD
def iniciar_processamento(url):

    if not url:
        return "URL inválida"

    thread = threading.Thread(
        target=processar,
        args=(url,)
    )

    thread.daemon = True

    thread.start()

    return "Processo iniciado"

# 📡 STATUS FRONTEND
def get_status():
    return status